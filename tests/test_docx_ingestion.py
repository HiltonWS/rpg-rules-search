from pathlib import Path

import fitz
from docx import Document

from rpg_rules_search.ingestion import docx as docx_ingestion


def test_docx_falls_back_to_searchable_pdf_without_libreoffice(
    tmp_path: Path, monkeypatch,
) -> None:
    source = tmp_path / "conditions.docx"
    destination = tmp_path / "conditions.pdf"
    document = Document()
    document.add_paragraph("Fascinado é uma condição mental.")
    table = document.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "Defesa"
    table.rows[0].cells[1].text = "18"
    document.save(source)
    monkeypatch.setattr(docx_ingestion.shutil, "which", lambda _command: None)

    docx_ingestion.convert_docx_to_pdf(source, destination)

    with fitz.open(destination) as pdf:
        text = "".join(page.get_text() for page in pdf)
    assert "Fascinado é uma condição mental." in text
    assert "Defesa | 18" in text