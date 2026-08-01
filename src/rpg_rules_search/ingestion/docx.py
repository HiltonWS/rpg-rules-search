from __future__ import annotations

import shutil
import subprocess
import textwrap
from pathlib import Path

import fitz
from docx import Document


class LibreOfficeUnavailableError(RuntimeError):
    pass


def _convert_docx_text_to_pdf(source: Path, destination: Path) -> None:
    source_document = Document(source)
    lines = [paragraph.text for paragraph in source_document.paragraphs]
    for table in source_document.tables:
        lines.extend(" | ".join(cell.text for cell in row.cells) for row in table.rows)

    pdf = fitz.open()
    page = pdf.new_page(width=595, height=842)
    y_position = 50.0
    for source_line in lines:
        wrapped_lines = textwrap.wrap(source_line, width=95) or [""]
        for line in wrapped_lines:
            if y_position > 792:
                page = pdf.new_page(width=595, height=842)
                y_position = 50.0
            page.insert_text((50, y_position), line, fontsize=10, fontname="helv")
            y_position += 14

    destination.parent.mkdir(parents=True, exist_ok=True)
    pdf.save(destination)
    pdf.close()


def convert_docx_to_pdf(source: Path, destination: Path) -> None:
    executable = shutil.which("libreoffice") or shutil.which("soffice")
    if executable is None:
        _convert_docx_text_to_pdf(source, destination)
        return

    destination.parent.mkdir(parents=True, exist_ok=True)
    result = subprocess.run(
        [
            executable,
            "--headless",
            "--convert-to",
            "pdf",
            "--outdir",
            str(destination.parent),
            str(source),
        ],
        capture_output=True,
        text=True,
        timeout=180,
        check=False,
    )
    generated = destination.parent / f"{source.stem}.pdf"
    if result.returncode != 0 or not generated.exists():
        detail = result.stderr.strip() or result.stdout.strip() or "conversão sem saída"
        raise RuntimeError(f"Falha ao converter DOCX: {detail}")
    generated.replace(destination)